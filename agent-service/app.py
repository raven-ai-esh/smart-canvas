from __future__ import annotations

from typing import Any
import contextvars
import copy
import json
import os
import time
import uuid
import html
import re
import mimetypes
import io
import pathlib
import base64
import hashlib
from urllib.parse import urlparse
import numpy as np
import pandas as pd
import contextlib

from contextlib import asynccontextmanager
from datetime import datetime, timedelta
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, Response
from openai import APIStatusError, AsyncOpenAI
from pydantic import BaseModel, ConfigDict, Field
import httpx
from mcp.client.session import ClientSession
from mcp.client.streamable_http import streamable_http_client
import tiktoken
from prometheus_client import Counter, Histogram, Gauge, CONTENT_TYPE_LATEST, generate_latest
from pythonjsonlogger import jsonlogger
from opentelemetry import trace
from opentelemetry.exporter.otlp.proto.http.trace_exporter import OTLPSpanExporter
from opentelemetry.instrumentation.fastapi import FastAPIInstrumentor
from opentelemetry.instrumentation.httpx import HTTPXClientInstrumentor
from opentelemetry.sdk.resources import Resource, SERVICE_NAME
from opentelemetry.sdk.trace import TracerProvider
from opentelemetry.sdk.trace.export import BatchSpanProcessor
from opentelemetry.util._once import Once

import logging
from bs4 import BeautifulSoup
import html2text
import fitz  # PyMuPDF
import openpyxl
from pypdf import PdfReader
import docx
from PIL import Image
try:
    from striprtf.striprtf import rtf_to_text
except Exception:  # pragma: no cover - optional dependency
    rtf_to_text = None

LOG_LEVEL = os.getenv("AGENT_LOG_LEVEL", os.getenv("LOG_LEVEL", "INFO")).upper()
LOG_TRUNCATE = int(os.getenv("AGENT_LOG_TRUNCATE", "2000"))
LOG_TRACE = os.getenv("LOG_TRACE", "false").lower() == "true"
METRICS_ENABLED = os.getenv("METRICS_ENABLED", "true").lower() != "false"
METRICS_PATH = os.getenv("METRICS_PATH", "/metrics")
OTEL_ENDPOINT = os.getenv("OTEL_EXPORTER_OTLP_ENDPOINT", "")
OTEL_SERVICE_NAME = os.getenv("OTEL_SERVICE_NAME", "smart-tracker-agent")
OTEL_LOG_LEVEL = os.getenv("OTEL_LOG_LEVEL", "")

request_id_ctx = contextvars.ContextVar("request_id", default=None)


class ContextFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        record.request_id = request_id_ctx.get() or ""
        if LOG_TRACE:
            span = trace.get_current_span()
            span_context = span.get_span_context() if span else None
            if span_context and span_context.trace_id:
                record.trace_id = f"{span_context.trace_id:032x}"
                record.span_id = f"{span_context.span_id:016x}"
            else:
                record.trace_id = ""
                record.span_id = ""
        return True


handler = logging.StreamHandler()
formatter = jsonlogger.JsonFormatter(
    "%(asctime)s %(levelname)s %(name)s %(message)s %(request_id)s %(trace_id)s %(span_id)s"
)
handler.setFormatter(formatter)
root_logger = logging.getLogger()
root_logger.handlers = [handler]
root_logger.setLevel(LOG_LEVEL)
root_logger.addFilter(ContextFilter())
if OTEL_LOG_LEVEL:
    logging.getLogger("opentelemetry").setLevel(OTEL_LOG_LEVEL)

app = FastAPI()

logger = logging.getLogger("agent")

_otel_once = Once()


def _init_tracing() -> None:
    if not OTEL_ENDPOINT:
        return
    def _setup() -> None:
        resource = Resource.create({SERVICE_NAME: OTEL_SERVICE_NAME})
        provider = TracerProvider(resource=resource)
        exporter = OTLPSpanExporter(endpoint=OTEL_ENDPOINT)
        provider.add_span_processor(BatchSpanProcessor(exporter))
        trace.set_tracer_provider(provider)
        FastAPIInstrumentor.instrument_app(app)
        HTTPXClientInstrumentor().instrument()
        logger.info("otel_ready", extra={"endpoint": OTEL_ENDPOINT, "service": OTEL_SERVICE_NAME})
    _otel_once.do_once(_setup)


REQUEST_COUNT = Counter(
    "http_requests_total",
    "Total HTTP requests",
    ["method", "path", "status"],
)
REQUEST_LATENCY = Histogram(
    "http_request_duration_seconds",
    "HTTP request duration in seconds",
    ["method", "path", "status"],
)
IN_FLIGHT = Gauge("http_in_flight_requests", "In-flight HTTP requests")


@app.middleware("http")
async def observability_middleware(request: Request, call_next):
    request_id = request.headers.get("x-request-id") or request.headers.get("x-correlation-id") or str(uuid.uuid4())
    token = request_id_ctx.set(request_id)
    start = time.time()
    if METRICS_ENABLED:
        IN_FLIGHT.inc()
    try:
        response = await call_next(request)
    except Exception as exc:
        duration = time.time() - start
        if METRICS_ENABLED:
            REQUEST_COUNT.labels(request.method, request.url.path, "500").inc()
            REQUEST_LATENCY.labels(request.method, request.url.path, "500").observe(duration)
            IN_FLIGHT.dec()
        logger.exception("http_error", extra={
            "method": request.method,
            "path": request.url.path,
            "status": 500,
            "duration_ms": int(duration * 1000),
        })
        request_id_ctx.reset(token)
        raise exc
    duration = time.time() - start
    if METRICS_ENABLED:
        REQUEST_COUNT.labels(request.method, request.url.path, str(response.status_code)).inc()
        REQUEST_LATENCY.labels(request.method, request.url.path, str(response.status_code)).observe(duration)
        IN_FLIGHT.dec()
    logger.info("http_request", extra={
        "method": request.method,
        "path": request.url.path,
        "status": response.status_code,
        "duration_ms": int(duration * 1000),
    })
    response.headers["x-request-id"] = request_id
    request_id_ctx.reset(token)
    return response


if METRICS_ENABLED:
    @app.get(METRICS_PATH)
    async def metrics() -> Response:
        return Response(generate_latest(), media_type=CONTENT_TYPE_LATEST)

_init_tracing()

AGENT_MODEL_CONTEXT_TOKENS = int(os.getenv("AGENT_MODEL_CONTEXT_TOKENS", "0"))
ASSISTANT_MODEL_CONTEXT_TOKENS = int(os.getenv("ASSISTANT_MODEL_CONTEXT_TOKENS", "0"))
PROMPT_PATH = os.getenv("AGENT_PROMPT_PATH", "/app/data/prompt.txt")
MODEL_CONTEXT_TOKENS = {
    "gpt-5.2": 400000,
}
FILE_BASE_URL = os.getenv("AGENT_FILE_BASE_URL", "http://api:8787")

DEFAULT_PROMPT_LINES = [
    "You are Raven, the Smart Tracker AI assistant.",
    "You can use MCP tools to read and update the canvas.",
    "Use tools when a user asks to inspect or change the canvas.",
    'Prefer node with action="create" for new cards and action="update" for edits.',
    "When creating edges between new cards, create the cards first and use their returned ids; do not use placeholder ids.",
    "get_state returns a summary by default (titles + metadata). Use node with action=\"read\" for full content when needed.",
    "Full get_state payloads are disabled; use node with action=\"read\" for full card details.",
    "If you only need a list of cards, use node with action=\"read\" and mode=\"summary\".",
    "Nodes have energy from 0 to 100 that represents the effort required to complete the card unless the user specifies otherwise.",
    "Energy propagates along edges from source nodes to target nodes.",
    "Each card has a base (own) energy you set directly; total card energy equals its base plus the sum of incoming energies, capped at 100%.",
    "List responses are capped; if a list is truncated, request specific items by id or use a smaller limit.",
    "Canvas participants are users who saved the canvas; only they can be tagged.",
    "Use MCP tool list_canvas_participants to fetch taggable people (id, name, email).",
    "Use MCP tool send_alert to notify a canvas participant via their enabled alerting channels. Pass userRef as the participant id (preferred) or their name/email/handle from list_canvas_participants.",
    "If you receive a message formatted like: \"\u041f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u0442\u0435\u043b\u044c <Name> \u043e\u0442\u0432\u0435\u0442\u0438\u043b \u043d\u0430 \u0441\u043e\u043e\u0431\u0449\u0435\u043d\u0438\u0435 \\\"...\\\" \u0442\u0435\u043a\u0441\u0442\u043e\u043c: \\\"...\\\"\", treat it as that user's reply to your earlier alert; it is not the current user relaying a message.",
    "When tagging someone in a card, include @Name in the content and update node.mentions with {id,label}.",
    "To tag everyone, include @all and add {id:\"all\", label:\"all\"} to node.mentions.",
    "For destructive actions (delete), ask for explicit confirmation first.",
    "If a tool fails, explain what happened and ask how to proceed.",
    "Keep responses concise and actionable.",
    "Current time: {{current_time}} | Date: {{current_date}} | Day: {{current_weekday}}",
]
DEFAULT_PROMPT = "\n".join(DEFAULT_PROMPT_LINES)

_prompt_cache: str | None = None
_prompt_mtime: float | None = None


def _detect_mime_from_name(name: str | None) -> str | None:
    if not name:
        return None
    mime, _ = mimetypes.guess_type(name)
    return mime


def _strip_html(text: str) -> str:
    soup = BeautifulSoup(text, "html.parser")
    for script in soup(["script", "style"]):
        script.decompose()
    return soup.get_text(separator=" ", strip=True)


def _html_to_text(text: str) -> str:
    try:
        h = html2text.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        h.ignore_emphasis = False
        h.body_width = 0
        return h.handle(text)
    except Exception:
        return _strip_html(text)


def _decode_bytes(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_text(content: bytes, mime: str | None, filename: str | None) -> str:
    mime_lower = (mime or "").lower()
    suffix = (pathlib.Path(filename).suffix.lower() if filename else "") if filename else ""

    if mime_lower in {"application/rtf", "text/rtf"} or suffix == ".rtf":
        text = _decode_bytes(content)
        if rtf_to_text:
            try:
                return rtf_to_text(text)
            except Exception:
                return text
        return text

    if mime_lower.startswith("text/") or suffix in {".txt", ".md", ".csv", ".log"}:
        return _decode_bytes(content)

    if mime_lower in {"text/markdown"} or suffix == ".md":
        return _decode_bytes(content)

    if mime_lower in {"text/html", "application/xhtml+xml"} or suffix in {".html", ".htm"}:
        return _html_to_text(_decode_bytes(content))

    if mime_lower == "application/pdf" or suffix == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    if mime_lower in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    } or suffix in {".docx", ".doc"}:
        try:
            document = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in document.paragraphs)
        except Exception:
            return ""

    if mime_lower.startswith("text/"):
        return _decode_bytes(content)

    return ""


def _extract_docx_text_and_tables(content: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception:
        return ""

    parts: list[str] = []
    paragraphs = [p.text for p in document.paragraphs if p.text]
    if paragraphs:
        parts.append("\n".join(paragraphs))

    table_blocks: list[str] = []
    for idx, table in enumerate(document.tables, 1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            table_blocks.append(f"[table {idx}]\n" + "\n".join(rows))
    if table_blocks:
        parts.append("\n\n".join(table_blocks))

    return "\n\n".join(parts)


def _extract_docx_images(content: bytes) -> list[bytes]:
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception:
        return []

    images: list[bytes] = []
    seen: set[str] = set()
    for rel in document.part.rels.values():
        try:
            if "image" not in rel.reltype:
                continue
            blob = rel.target_part.blob
            digest = hashlib.sha1(blob).hexdigest()
            if digest in seen:
                continue
            seen.add(digest)
            images.append(blob)
        except Exception:
            continue
    return images


def _image_bytes_to_png_b64(data: bytes) -> str | None:
    try:
        img = Image.open(io.BytesIO(data))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    except Exception:
        return None


def _should_use_vision_for_docx(text: str, question: str) -> bool:
    if not text or len(text.strip()) < 200:
        return True
    q = (question or "").lower()
    keywords = [
        "image", "figure", "diagram", "chart", "graph", "picture", "photo",
        "\u0440\u0438\u0441\u0443\u043d\u043e\u043a", "\u0434\u0438\u0430\u0433\u0440\u0430\u043c", "\u0433\u0440\u0430\u0444\u0438\u043a", "\u0441\u0445\u0435\u043c",
    ]
    return any(k in q for k in keywords)


async def _vision_answer_image(
    client: AsyncOpenAI,
    model: str | None,
    question: str,
    image_bytes: bytes,
    image_index: int,
) -> dict[str, Any]:
    b64 = _image_bytes_to_png_b64(image_bytes)
    if not b64:
        return {"image": image_index + 1, "answer": ""}
    content = [
        {
            "type": "text",
            "text": (
                "Answer the question using only the provided image. "
                "Return concise text; if nothing relevant, say 'not found'. "
            ),
        },
        {
            "type": "text",
            "text": f"Question: {question}\nImage: {image_index + 1}",
        },
        {
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{b64}"},
        },
    ]
    chat = await client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": content}],
        temperature=0,
    )
    answer = chat.choices[0].message.content or ""
    return {"image": image_index + 1, "answer": answer.strip()}


def _load_tables(content: bytes, mime: str | None, filename: str | None) -> dict[str, pd.DataFrame]:
    """Load tabular file into dict of DataFrames keyed by sheet name."""
    mime_lower = (mime or "").lower()
    suffix = (pathlib.Path(filename).suffix.lower() if filename else "")
    dfs: dict[str, pd.DataFrame] = {}
    buf = io.BytesIO(content)
    logger.info(
        "table_load_start",
        extra={
            "mime": mime_lower,
            "suffix": suffix,
            "size_bytes": len(content),
        },
    )

    try:
        if suffix in {".csv", ".tsv"} or mime_lower in {"text/csv", "text/tab-separated-values"}:
            sep = "\t" if suffix == ".tsv" or "tab" in mime_lower else None
            csv_encodings = [None, "utf-8", "utf-8-sig", "cp1251", "latin-1"]
            last_error: Exception | None = None
            for enc in csv_encodings:
                try:
                    buf.seek(0)
                    df = pd.read_csv(
                        buf,
                        sep=sep,
                        engine="python",
                        encoding=enc if enc else None,
                    )
                    logger.info(
                        "table_load_csv_ok",
                        extra={
                            "rows": df.shape[0],
                            "cols": df.shape[1],
                            "encoding": enc or "default",
                        },
                    )
                    dfs["Sheet1"] = df
                    return dfs
                except Exception as exc:
                    last_error = exc
                    continue
            # last-resort: let pandas sniff sep/encoding with errors replaced
            try:
                buf.seek(0)
                df = pd.read_csv(buf, sep=None, engine="python", encoding_errors="replace")
                logger.info(
                    "table_load_csv_ok_fallback",
                    extra={
                        "rows": df.shape[0],
                        "cols": df.shape[1],
                        "error": str(last_error) if last_error else "",
                    },
                )
                dfs["Sheet1"] = df
                return dfs
            except Exception as exc2:
                logger.error(
                    "table_load_csv_failed",
                    extra={"error": str(exc2), "mime": mime_lower, "suffix": suffix},
                )
                return {}

        if suffix in {".xlsx", ".xls", ".ods"} or mime_lower in {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
            "application/vnd.oasis.opendocument.spreadsheet",
            }:
            try:
                dfs_raw = pd.read_excel(buf, sheet_name=None, dtype_backend="pyarrow")
                logger.info(
                    "table_load_excel_ok",
                    extra={"sheets": list(dfs_raw.keys())},
                )
            except Exception as exc:
                buf.seek(0)
                try:
                    dfs_raw = pd.read_excel(buf, sheet_name=None, engine="openpyxl")
                    logger.info(
                        "table_load_excel_ok_openpyxl",
                        extra={"sheets": list(dfs_raw.keys()), "error": str(exc)},
                    )
                except Exception:
                    dfs_raw = {}
            return {str(k): v for k, v in dfs_raw.items()}

        if suffix in {".parquet"} or mime_lower == "application/octet-stream":
            buf.seek(0)
            try:
                df = pd.read_parquet(buf)
                logger.info(
                    "table_load_parquet_ok",
                    extra={"rows": df.shape[0], "cols": df.shape[1]},
                )
                dfs["Sheet1"] = df
                return dfs
            except Exception as exc:
                logger.error(
                    "table_load_parquet_failed",
                    extra={"error": str(exc), "mime": mime_lower, "suffix": suffix},
                )
                return {}
    except Exception as exc:
        logger.warning("table_load_failed", extra={"error": str(exc), "mime": mime, "suffix": suffix})
        # last-resort: openpyxl manual read
        try:
            buf.seek(0)
            wb = openpyxl.load_workbook(buf, data_only=True, read_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                header = rows[0]
                data = rows[1:] if len(rows) > 1 else []
                df = pd.DataFrame(data, columns=header)
                dfs[str(sheet)] = df
            wb.close()
            if dfs:
                logger.info(
                    "table_load_openpyxl_manual_ok",
                    extra={"sheets": list(dfs.keys())},
                )
            return dfs
        except Exception as exc2:
            logger.error("table_load_failed_final", extra={"error": str(exc2), "mime": mime, "suffix": suffix})
            return {}
    return {}


def _df_sample(df: pd.DataFrame, rows: int = 5) -> str:
    try:
        return df.head(rows).to_markdown(index=False)
    except Exception:
        return df.head(rows).to_string(index=False)

def _format_table_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, pd.DataFrame):
        try:
            return value.head(20).to_markdown(index=False)
        except Exception:
            return value.head(20).to_string(index=False)
    if isinstance(value, pd.Series):
        try:
            return value.head(20).to_string(index=False)
        except Exception:
            return str(value.head(20))
    if isinstance(value, (list, dict, tuple)):
        try:
            return json.dumps(value, ensure_ascii=False)
        except Exception:
            return str(value)
    try:
        # numpy scalars
        if hasattr(value, "item"):
            return str(value.item())
    except Exception:
        pass
    return str(value)

def _extract_json_payload(text: str) -> dict[str, Any] | None:
    if not text:
        return None
    match = re.search(r"```json\\s*(\\{.*?\\})```", text, re.S | re.I)
    if not match:
        match = re.search(r"```\\s*(\\{.*?\\})```", text, re.S)
    payload = None
    if match:
        payload = match.group(1).strip()
    else:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            payload = text[start:end + 1]
    if not payload:
        return None
    try:
        return json.loads(payload)
    except Exception:
        return None


def _safe_table_exec(code: str, dfs: dict[str, pd.DataFrame]) -> dict[str, Any]:
    allowed_builtin_names = {
        "__import__",
        "print",
        "len",
        "range",
        "int",
        "float",
        "str",
        "bool",
        "list",
        "dict",
        "set",
        "tuple",
        "min",
        "max",
        "sum",
        "abs",
        "sorted",
        "enumerate",
        "any",
        "all",
        "zip",
        "round",
        "isinstance",
    }
    if isinstance(__builtins__, dict):
        builtins_map = {name: __builtins__[name] for name in allowed_builtin_names if name in __builtins__}
    else:
        builtins_map = {name: getattr(__builtins__, name) for name in allowed_builtin_names}
    allowed_builtins = {"__builtins__": builtins_map}
    env = {"dfs": dfs, "pd": pd, "np": np}
    stdout_buf = io.StringIO()
    try:
        compiled = compile(code, "<table_exec>", "exec")
        with contextlib.redirect_stdout(stdout_buf):
            exec(compiled, allowed_builtins, env)
        return {"ok": True, "stdout": stdout_buf.getvalue().strip(), "result": env.get("result")}
    except Exception as exc:
        error_text = str(exc)
        error_type = "execution_failed"
        missing_name = None
        if isinstance(exc, NameError):
            match = re.search(r"name '([^']+)' is not defined", error_text)
            if match:
                missing_name = match.group(1)
                # Only flag as restricted when it's a real builtin that's blocked.
                try:
                    builtin_names = set(dir(__builtins__))
                except Exception:
                    builtin_names = set()
                if missing_name in builtin_names and missing_name not in allowed_builtin_names:
                    error_type = "restricted_builtin"
        return {
            "ok": False,
            "error": error_text,
            "error_type": error_type,
            "missing_name": missing_name,
            "stdout": stdout_buf.getvalue().strip(),
            "result": None,
        }


def _pdf_page_texts(content: bytes) -> list[str]:
    """Extract plain text per page using PyMuPDF; fallback to empty strings."""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception:
        return []
    texts: list[str] = []
    for page in doc:
        try:
            txt = page.get_text("text") or ""
        except Exception:
            txt = ""
        texts.append(txt)
    doc.close()
    return texts


def _render_pdf_page_png(content: bytes, page_index: int, target_width: int = 1400) -> bytes:
    """Render a single PDF page to PNG bytes."""
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        page = doc[page_index]
        zoom = max(1.0, min(3.0, target_width / max(1.0, page.rect.width)))
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return pix.tobytes("png")
    finally:
        doc.close()


def _is_bad_page_text(text: str) -> bool:
    if not text or len(text.strip()) < 80:
        return True
    bad_chars = text.count("\ufffd")
    non_alpha = sum(1 for ch in text if not ch.isalpha() and not ch.isspace())
    total = len(text) or 1
    return bad_chars > 0 or (non_alpha / total) > 0.6


def _count_tokens(text: str, model: str | None) -> int:
    enc_name = None
    try:
        if model:
            enc = tiktoken.encoding_for_model(model)
        else:
            enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except Exception:
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))


def _chunk_text(text: str, model: str | None, chunk_tokens: int = 900, overlap_tokens: int = 200) -> list[str]:
    try:
        enc = tiktoken.encoding_for_model(model) if model else tiktoken.get_encoding("cl100k_base")
    except Exception:
        enc = tiktoken.get_encoding("cl100k_base")
    tokens = enc.encode(text)
    chunks = []
    i = 0
    while i < len(tokens):
        window = tokens[i : i + chunk_tokens]
        chunks.append(enc.decode(window))
        i += chunk_tokens - overlap_tokens
    return chunks


async def _vision_answer_page(
    client: AsyncOpenAI,
    model: str | None,
    question: str,
    page_png: bytes,
    page_index: int,
) -> dict[str, Any]:
    b64 = base64.b64encode(page_png).decode("ascii")
    content = [
        {
            "type": "text",
            "text": (
                "You will see a PDF page image. "
                "Answer the user question using ONLY this page. "
                "Return concise text; if nothing relevant, say 'not found'. "
            ),
        },
        {
            "type": "text",
            "text": f"Question: {question}\nPage: {page_index + 1}",
        },
        {
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{b64}"},
        },
    ]
    chat = await client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": content}],
        temperature=0,
    )
    answer = chat.choices[0].message.content or ""
    return {"page": page_index + 1, "answer": answer.strip()}


async def _run_table_agent(
    dfs: dict[str, pd.DataFrame],
    question: str,
    model: str | None,
    client: AsyncOpenAI,
    max_attempts: int = 4,
) -> dict[str, Any]:
    sheet_summaries = []
    for name, df in dfs.items():
        rows, cols = df.shape
        summary = {
            "sheet": name,
            "rows": rows,
            "cols": cols,
            "columns": list(df.columns[:20]),
            "preview": _df_sample(df),
        }
        sheet_summaries.append(summary)

    allowed_builtin_list = [
        "print", "len", "range", "int", "float", "str", "bool", "list", "dict", "set", "tuple",
        "min", "max", "sum", "abs", "sorted", "enumerate", "any", "all", "zip", "round", "isinstance",
    ]
    sys_prompt = (
        "You are a data analyst. You can run Python pandas code to answer the user's question.\n"
        "- Available DataFrames are in dict `dfs`, keyed by sheet name.\n"
        "- Choose the right sheet by name (e.g., dfs['Sheet1']).\n"
        "- Use only pandas/numpy; no file/network/unsafe operations.\n"
        "- Return ONLY JSON (no markdown) with fields: action, code, answer.\n"
        "- action: \"run_code\" to execute code, or \"answer\" if you can answer now.\n"
        "- When action is \"run_code\", code MUST assign `result = ...` (JSON-serializable).\n"
        f"- Allowed builtins: {', '.join(allowed_builtin_list)}.\n"
        "- Keep code short and efficient.\n"
    )

    sheets_text = "\n\n".join(
        f"Sheet: {s['sheet']} (rows={s['rows']}, cols={s['cols']})\n"
        f"Columns: {s['columns']}\nSample:\n{s['preview']}"
        for s in sheet_summaries
    )

    messages = [
        {"role": "system", "content": sys_prompt},
        {
            "role": "user",
            "content": (
                f"Question: {question}\n\n"
                "Tables available:\n"
                f"{sheets_text}\n\n"
                "Reply ONLY with JSON as specified in the system instructions."
            ),
        },
    ]

    last_error = None
    used_sheets: list[str] = []
    for attempt in range(1, max_attempts + 1):
        logger.info("table_agent_attempt", extra={"attempt": attempt})
        chat = await client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0,
        )
        reply = chat.choices[0].message.content or ""
        payload = _extract_json_payload(reply)
        if not payload:
            last_error = "invalid_json"
            messages.append({"role": "assistant", "content": reply})
            messages.append(
                {
                    "role": "user",
                    "content": "Your response was not valid JSON. Please respond with JSON only.",
                }
            )
            continue

        action = str(payload.get("action") or "").strip().lower()
        code = str(payload.get("code") or "").strip()
        answer = str(payload.get("answer") or "").strip()
        if action == "answer" and answer:
            return {
                "answer": answer,
                "mode": "table_agent",
                "used_chunks": [{"sheet": s} for s in used_sheets],
            }
        if action != "run_code" or not code:
            last_error = "missing_code"
            messages.append({"role": "assistant", "content": reply})
            messages.append(
                {
                    "role": "user",
                    "content": "You must respond with action=run_code and provide Python code that assigns `result`.",
                }
            )
            continue

        # crude sheet detection
        for name in dfs.keys():
            if f"dfs['{name}'" in code or f'dfs["{name}"' in code:
                used_sheets.append(name)

        exec_result = _safe_table_exec(code, dfs)
        logger.info(
            "table_agent_exec",
            extra={
                "attempt": attempt,
                "used_sheets": used_sheets,
                "ok": exec_result.get("ok"),
                "error": exec_result.get("error"),
                "error_type": exec_result.get("error_type"),
                "missing_name": exec_result.get("missing_name"),
                "stdout_preview": (exec_result.get("stdout") or "")[:200],
                "result_preview": _format_table_value(exec_result.get("result"))[:200],
                "code_preview": code[:200],
            },
        )
        if exec_result.get("ok"):
            stdout = exec_result.get("stdout", "")
            result_text = _format_table_value(exec_result.get("result"))
            combined = result_text or stdout
            if combined:
                messages.append({"role": "assistant", "content": reply})
                messages.append(
                    {
                        "role": "user",
                        "content": (
                            "Execution result:\n"
                            f"{combined}\n\n"
                            "If you can answer now, respond with action=answer and a concise answer. "
                            "Otherwise, run more code."
                        ),
                    }
                )
                continue
            last_error = "empty_output"
            messages.append({"role": "assistant", "content": reply})
            messages.append(
                {
                    "role": "user",
                    "content": "Your code produced no output. Please assign result = ... (JSON-serializable).",
                }
            )
            continue

        last_error = exec_result.get("error", "execution_failed")
        error_type = exec_result.get("error_type")
        missing_name = exec_result.get("missing_name")
        messages.append(
            {
                "role": "assistant",
                "content": reply,
            }
        )
        if error_type == "restricted_builtin" and missing_name:
            messages.append(
                {
                    "role": "user",
                    "content": (
                        f"Your code used a name that is not available in the sandbox: '{missing_name}'. "
                        f"Allowed builtins: {', '.join(allowed_builtin_list)}. "
                        "Use only dfs/pd/np and allowed builtins, then respond with JSON only."
                    ),
                }
            )
        elif missing_name:
            messages.append(
                {
                    "role": "user",
                    "content": (
                        f"Your code failed because '{missing_name}' was used before it was defined. "
                        "Define variables before use (e.g., `rows = []`) and respond with JSON only."
                    ),
                }
            )
        else:
            messages.append(
                {
                    "role": "user",
                    "content": f"Your code failed with error: {last_error}. Please fix and respond with JSON only.",
                }
            )

    return {
        "answer": "",
        "mode": "table_error",
        "error": "table_exec_failed",
        "used_chunks": [{"sheet": s} for s in used_sheets],
    }


def _is_internal_file_url(raw_url: Any, resolved_url: str | None) -> bool:
    if isinstance(raw_url, str) and raw_url.startswith("/"):
        return True
    if not resolved_url:
        return False
    base = urlparse(FILE_BASE_URL)
    target = urlparse(resolved_url)
    if not base.netloc or not target.netloc:
        return False
    return base.netloc == target.netloc


async def _run_doc_search(
    args: dict[str, Any],
    model: str | None,
    client: AsyncOpenAI,
    mcp_token: str | None = None,
) -> dict[str, Any]:
    raw_url = args.get("download_url") or args.get("url")
    if raw_url and isinstance(raw_url, str) and raw_url.startswith("/"):
        download_url = FILE_BASE_URL.rstrip("/") + raw_url
    else:
        download_url = raw_url
    search_request = args.get("search_request") or args.get("query") or ""
    if not download_url or not search_request:
        raise HTTPException(status_code=400, detail="download_url and search_request are required")

    request_headers: dict[str, str] = {}
    if isinstance(mcp_token, str) and mcp_token.strip() and _is_internal_file_url(raw_url, download_url):
        request_headers["Authorization"] = f"Bearer {mcp_token.strip()}"

    async with httpx.AsyncClient(timeout=30) as http_client:
        resp = await http_client.get(download_url, headers=request_headers or None)
        resp.raise_for_status()
        content = resp.content
        mime = resp.headers.get("content-type", "").split(";")[0].strip() or None
        filename = (
            args.get("file_name")
            or args.get("filename")
            or args.get("fileName")
        )
        if not filename:
            # fallback to URL path segment
            filename = download_url.split("?")[0].split("/")[-1] if download_url else None
        if not filename:
            filename = _mask_secret(download_url.split("/")[-1]) if download_url else ""
        if not mime:
            mime = _detect_mime_from_name(filename)

    suffix = pathlib.Path(filename).suffix.lower() if filename else ""

    logger.info(
        "doc_search_file",
        extra={
            "file": filename,
            "mime": mime,
            "suffix": suffix,
            "size_bytes": len(content),
            "download_url": _mask_secret(download_url),
        },
    )

    # Table docs branch
    table_suffixes = {".csv", ".tsv", ".xlsx", ".xls", ".ods", ".parquet"}
    table_mimes = {
        "text/csv",
        "text/tab-separated-values",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "application/vnd.oasis.opendocument.spreadsheet",
    }
    if suffix in table_suffixes or (mime and mime.lower() in table_mimes):
        dfs = _load_tables(content, mime, filename)
        if not dfs:
            return {
                "answer": "table_load_failed",
                "mode": "table_error",
                "error": "failed_to_load_table",
                "token_count": 0,
                "file": {"name": filename, "mime": mime, "size_bytes": len(content)},
                "used_chunks": [],
            }
        logger.info(
            "table_loaded",
            extra={
                "file": filename,
                "mime": mime,
                "size": len(content),
                "sheets": {k: list(v.shape) for k, v in dfs.items()},
            },
        )
        table_result = await _run_table_agent(dfs, search_request, model, client)
        return {
            "answer": table_result.get("answer", ""),
            "mode": table_result.get("mode", "table_agent"),
            "token_count": 0,
            "file": {"name": filename, "mime": mime, "size_bytes": len(content)},
            "used_chunks": table_result.get("used_chunks", []),
            "error": table_result.get("error"),
        }

    image_suffixes = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"}
    if (mime and mime.lower().startswith("image/")) or (suffix in image_suffixes):
        logger.info(
            "doc_search_mode",
            extra={
                "mode": "image",
                "mime": mime,
                "suffix": suffix,
            },
        )
        vision = await _vision_answer_image(client, model, search_request, content, 0)
        return {
            "answer": vision.get("answer", ""),
            "mode": "vision_image",
            "token_count": 0,
            "file": {"name": filename, "mime": mime, "size_bytes": len(content)},
            "used_chunks": [{"image": vision.get("image")}],
        }

    logger.info(
        "doc_search_mode",
        extra={
            "mode": "text",
            "mime": mime,
            "suffix": suffix,
        },
    )

    vision_pages: list[dict[str, Any]] = []
    combined_text = ""

    if (mime and mime.lower() in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }) or (suffix in {".docx", ".doc"}):
        combined_text = _extract_docx_text_and_tables(content)
        images = _extract_docx_images(content) if suffix == ".docx" or (mime and "wordprocessingml" in mime.lower()) else []
        if images and _should_use_vision_for_docx(combined_text, search_request):
            MAX_VISION_IMAGES = 3
            for idx, image_bytes in enumerate(images[:MAX_VISION_IMAGES]):
                try:
                    vis = await _vision_answer_image(client, model, search_request, image_bytes, idx)
                    if vis.get("answer"):
                        vision_pages.append({"image": vis["image"], "preview": vis["answer"][:200]})
                except Exception:
                    continue
            if vision_pages:
                vision_text = "\n\n".join(
                    f"[vision image {item['image']}]\n{item['preview']}" for item in vision_pages if item.get("preview")
                )
                combined_parts = [part for part in [combined_text, vision_text] if part]
                combined_text = "\n\n".join(combined_parts) if combined_parts else combined_text
    elif suffix == ".rtf" or (mime and mime.lower() in {"application/rtf", "text/rtf"}):
        combined_text = _extract_text(content, mime, filename)
    else:
        combined_text = _extract_text(content, mime, filename)

    if (mime and mime.lower() == "application/pdf") or suffix == ".pdf":
        page_texts = _pdf_page_texts(content)
        good_pages_text = "\n\n".join(
            f"[page {idx + 1}]\n{txt}"
            for idx, txt in enumerate(page_texts)
            if txt and txt.strip()
        )

        # Pick pages with bad/empty text for vision Q&A
        problem_pages = [idx for idx, txt in enumerate(page_texts) if _is_bad_page_text(txt)]
        MAX_VISION_PAGES = 3
        for idx in problem_pages[:MAX_VISION_PAGES]:
            try:
                png = _render_pdf_page_png(content, idx)
                vis = await _vision_answer_page(client, model, search_request, png, idx)
                if vis.get("answer"):
                    vision_pages.append({"page": vis["page"], "preview": vis["answer"][:200]})
            except Exception:
                continue

        vision_text = "\n\n".join(
            f"[vision page {item['page']}]\n{item['preview']}" for item in vision_pages if item.get("preview")
        )

        combined_parts = [part for part in [good_pages_text, vision_text] if part]
        combined_text = "\n\n".join(combined_parts) if combined_parts else ""

    token_count = _count_tokens(combined_text, model)

    if token_count == 0:
        return {
            "answer": "",
            "mode": "empty",
            "token_count": 0,
            "file": {"name": filename, "mime": mime, "size_bytes": len(content)},
            "used_chunks": [],
        }

    DIRECT_THRESHOLD = 100_000
    if token_count <= DIRECT_THRESHOLD:
        prompt = [
            {
                "role": "system",
                "content": "You are Raven. Answer the user's request using ONLY the provided document text.",
            },
            {
                "role": "user",
                "content": f"Request: {search_request}\n\nDocument:\n{combined_text}",
            },
        ]
        chat = await client.chat.completions.create(
            model=model,
            messages=prompt,
            temperature=0,
        )
        answer = chat.choices[0].message.content or ""
        return {
            "answer": answer,
            "mode": "direct",
            "token_count": token_count,
            "file": {"name": filename, "mime": mime, "size_bytes": len(content)},
            "used_chunks": vision_pages,
        }

    # RAG path
    chunks = _chunk_text(combined_text, model, chunk_tokens=900, overlap_tokens=200)
    embeddings = await client.embeddings.create(
        model="text-embedding-3-large",
        input=chunks,
    )
    query_emb = await client.embeddings.create(
        model="text-embedding-3-large",
        input=search_request,
    )
    q = np.array(embeddings.data[-1].embedding if False else query_emb.data[0].embedding)
    scores = []
    for idx, item in enumerate(embeddings.data):
        v = np.array(item.embedding)
        score = float(np.dot(q, v) / (np.linalg.norm(q) * np.linalg.norm(v) + 1e-8))
        scores.append((score, idx))
    scores.sort(reverse=True)
    top_k = 6
    chosen = scores[:top_k]
    top_chunks = []
    used = []
    for rank, (score, idx) in enumerate(chosen, 1):
        chunk_text = chunks[idx]
        top_chunks.append(chunk_text)
        used.append(
            {
                "rank": rank,
                "score": round(score, 4),
                "preview": chunk_text[:240],
            }
        )

    rag_prompt = [
        {
            "role": "system",
            "content": "You are Raven. Answer the user's request using ONLY the provided document excerpts. If not enough info, say so.",
        },
        {
            "role": "user",
            "content": f"Request: {search_request}\n\nRelevant excerpts:\n" + "\n\n---\n\n".join(top_chunks),
        },
    ]
    chat = await client.chat.completions.create(
        model=model,
        messages=rag_prompt,
        temperature=0,
    )
    answer = chat.choices[0].message.content or ""
    return {
        "answer": answer,
        "mode": "rag",
        "token_count": token_count,
        "file": {"name": filename, "mime": mime, "size_bytes": len(content)},
        "used_chunks": used + vision_pages,
    }


async def _safe_doc_search(
    args: dict[str, Any],
    model: str | None,
    client: AsyncOpenAI,
    mcp_token: str | None = None,
) -> dict[str, Any]:
    try:
        result = await _run_doc_search(args, model, client, mcp_token=mcp_token)
        return {"ok": True, "result": result}
    except Exception as exc:
        logger.exception("doc_search_error", extra={"error": str(exc)})
        return {"ok": False, "error": str(exc)}


class MCPConfig(BaseModel):
    model_config = ConfigDict(extra="allow")

    url: str | None = None
    token: str | None = None
    sessionId: str | None = None
    userId: str | None = None
    allowedTools: list[str] | None = None


class AgentRunRequest(BaseModel):
    model_config = ConfigDict(extra="allow")

    apiKey: str = Field(..., min_length=1)
    model: str = Field(..., min_length=1)
    instructions: str | None = None
    userName: str | None = None
    input: str | list[dict[str, Any]]
    maxTurns: int | None = None
    temperature: float | None = 0.3
    openaiBaseUrl: str | None = None
    openaiTimeoutMs: int | None = None
    webSearchEnabled: bool | None = None
    mcp: MCPConfig | None = None


class AgentRunResponse(BaseModel):
    output: str
    lastResponseId: str | None = None
    context: dict[str, Any] | None = None
    trace: dict[str, Any] | None = None


class AgentContextRequest(BaseModel):
    model_config = ConfigDict(extra="allow")

    model: str = Field(..., min_length=1)
    instructions: str | None = None
    input: Any = None
    userName: str | None = None


class AgentContextResponse(BaseModel):
    context: dict[str, Any]


class AssistantResponse(BaseModel):
    message: str
    reasoning: str | None = None

class PromptResponse(BaseModel):
    prompt: str

class PromptUpdateRequest(BaseModel):
    prompt: str = Field(..., min_length=1)


def _format_output(value: Any) -> str:
    if isinstance(value, str):
        return value
    return str(value) if value is not None else ""

def _mask_secret(value: Any, keep: int = 4) -> str:
    raw = value if isinstance(value, str) else str(value) if value is not None else ""
    if not raw:
        return ""
    if len(raw) <= keep * 2:
        return f"{raw[:1]}...{raw[-1:]}"
    return f"{raw[:keep]}...{raw[-keep:]}"

def _summarize_output_items(output: Any) -> dict[str, int]:
    counts: dict[str, int] = {}
    if not isinstance(output, list):
        return counts
    for item in output:
        item_type = item.get("type") if isinstance(item, dict) else getattr(item, "type", None)
        label = item_type if isinstance(item_type, str) and item_type else "unknown"
        counts[label] = counts.get(label, 0) + 1
    return counts


def _extract_response_text(response: Any) -> str:
    if hasattr(response, "output_text") and isinstance(response.output_text, str):
        return response.output_text
    output = getattr(response, "output", None)
    if not isinstance(output, list):
        return ""
    for item in output:
        content = item.get("content") if isinstance(item, dict) else getattr(item, "content", None)
        if not isinstance(content, list):
            continue
        for part in content:
            part_type = part.get("type") if isinstance(part, dict) else getattr(part, "type", None)
            text = part.get("text") if isinstance(part, dict) else getattr(part, "text", None)
            if part_type == "output_text" and isinstance(text, str):
                return text
    return ""

def _ensure_prompt_file() -> str:
    path = PROMPT_PATH
    parent = os.path.dirname(path)
    if parent and not os.path.exists(parent):
        os.makedirs(parent, exist_ok=True)
    if not os.path.exists(path):
        with open(path, "w", encoding="utf-8") as handle:
            handle.write(DEFAULT_PROMPT + "\n")
    return path

def _load_prompt_text() -> str:
    path = _ensure_prompt_file()
    global _prompt_cache, _prompt_mtime
    try:
        mtime = os.path.getmtime(path)
    except OSError:
        mtime = None
    if _prompt_cache is not None and mtime is not None and _prompt_mtime == mtime:
        return _prompt_cache
    try:
        with open(path, "r", encoding="utf-8") as handle:
            text = handle.read()
    except OSError:
        text = DEFAULT_PROMPT
    text = text.strip()
    if not text:
        text = DEFAULT_PROMPT
    _prompt_cache = text
    _prompt_mtime = mtime
    return text

def _save_prompt_text(value: str) -> str:
    text = value.strip()
    if not text:
        raise ValueError("prompt_required")
    path = _ensure_prompt_file()
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(text + "\n")
    global _prompt_cache, _prompt_mtime
    _prompt_cache = text
    try:
        _prompt_mtime = os.path.getmtime(path)
    except OSError:
        _prompt_mtime = None
    return text

def _build_instructions(user_name: str | None, extra: str | None) -> str:
    template = _load_prompt_text()
    time_line = "Current time: {{current_time}} | Date: {{current_date}} | Day: {{current_weekday}}"
    if not any(token in template for token in ("{{current_time}}", "{{current_date}}", "{{current_weekday}}", "{{current_datetime}}")):
        template = "\n".join([template, time_line]).strip()
    now = datetime.now()
    current_time = now.strftime("%H:%M")
    current_date = now.strftime("%Y-%m-%d")
    current_weekday = now.strftime("%A")
    rendered = (
        template
        .replace("{{current_time}}", current_time)
        .replace("{{current_date}}", current_date)
        .replace("{{current_weekday}}", current_weekday)
        .replace("{{current_datetime}}", f"{current_date} {current_time} ({current_weekday})")
    )
    parts = [rendered]
    if isinstance(user_name, str) and user_name.strip():
        parts.append(f'The user name is "{user_name.strip()}".')
    if isinstance(extra, str) and extra.strip():
        parts.append(extra.strip())
    return "\n".join(parts)

def _parse_tool_args(raw: Any) -> dict[str, Any]:
    if not raw:
        return {}
    if isinstance(raw, dict):
        return raw
    if isinstance(raw, str):
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return {}
    return {}

def _normalize_tool_schema(schema: Any) -> dict[str, Any]:
    if isinstance(schema, dict):
        normalized = copy.deepcopy(schema)
    else:
        normalized = {}
    if normalized.get("type") is None:
        normalized["type"] = "object"
    if not isinstance(normalized.get("properties"), dict):
        normalized["properties"] = {}
    for key, prop in normalized["properties"].items():
        if not isinstance(prop, dict):
            normalized["properties"][key] = {"type": "string"}
            continue
        if not any(k in prop for k in ("type", "anyOf", "oneOf", "allOf")):
            prop["type"] = "object"
        if prop.get("type") == "object":
            if not isinstance(prop.get("properties"), dict):
                prop["properties"] = {}
            prop["additionalProperties"] = False
    normalized["additionalProperties"] = False
    return normalized

def _serialize_tool_result(result: Any) -> str:
    if result is None:
        return ""
    structured = getattr(result, "structuredContent", None)
    if structured is not None:
        return json.dumps(structured, ensure_ascii=False)
    content = getattr(result, "content", None)
    if isinstance(content, list):
        payload = []
        for block in content:
            if hasattr(block, "model_dump"):
                payload.append(block.model_dump())
            else:
                payload.append(block)
        return json.dumps(payload, ensure_ascii=False)
    return json.dumps(str(result), ensure_ascii=False)

def _extract_function_calls(output: Any) -> list[Any]:
    calls: list[Any] = []
    if not isinstance(output, list):
        return calls
    for item in output:
        item_type = item.get("type") if isinstance(item, dict) else getattr(item, "type", None)
        if item_type == "function_call":
            calls.append(item)
    return calls

def _extract_web_search_items(output: Any) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    if not isinstance(output, list):
        return items
    for item in output:
        item_type = item.get("type") if isinstance(item, dict) else getattr(item, "type", None)
        if not isinstance(item_type, str) or "web_search" not in item_type:
            continue
        if isinstance(item, dict):
            items.append(item)
            continue
        if hasattr(item, "model_dump"):
            items.append(item.model_dump())
            continue
        items.append({"type": item_type, "value": str(item)})
    return items

def _tool_call_priority(call: Any) -> int:
    name = call.get("name") if isinstance(call, dict) else getattr(call, "name", None)
    args_raw = call.get("arguments") if isinstance(call, dict) else getattr(call, "arguments", None)
    args = _parse_tool_args(args_raw)
    action = args.get("action") if isinstance(args, dict) else None
    if name == "edge" and action == "create":
        return 10
    return 0

def _prioritize_tool_calls(calls: list[Any]) -> list[Any]:
    indexed = list(enumerate(calls))
    indexed.sort(key=lambda pair: (_tool_call_priority(pair[1]), pair[0]))
    return [call for _, call in indexed]

@asynccontextmanager
async def mcp_session_context(mcp_config: MCPConfig | None, timeout_s: float | None):
    if not mcp_config or not mcp_config.url:
        yield None, []
        return
    headers: dict[str, str] = {}
    if mcp_config.token:
        headers["authorization"] = f"Bearer {mcp_config.token}"
    if mcp_config.sessionId:
        headers["x-session-id"] = mcp_config.sessionId
    if mcp_config.userId:
        headers["x-user-id"] = mcp_config.userId

    http_client = httpx.AsyncClient(headers=headers, timeout=timeout_s)
    async with http_client:
        async with streamable_http_client(mcp_config.url, http_client=http_client, terminate_on_close=True) as streams:
            read_stream, write_stream, _get_session_id = streams
            read_timeout = timedelta(seconds=timeout_s) if timeout_s else None
            async with ClientSession(read_stream, write_stream, read_timeout_seconds=read_timeout) as session:
                await session.initialize()
                tools_result = await session.list_tools()
                yield session, tools_result.tools

def _estimate_size(value: Any) -> int:
    if value is None:
        return 0
    if isinstance(value, str):
        return len(value)
    try:
        return len(json.dumps(value, ensure_ascii=False))
    except Exception:
        return len(str(value))

def _safe_log_payload(value: Any, max_chars: int | None = None) -> str:
    limit = max_chars if isinstance(max_chars, int) and max_chars > 0 else LOG_TRUNCATE
    if isinstance(value, str):
        text = value
    else:
        try:
            text = json.dumps(value, ensure_ascii=False)
        except Exception:
            text = str(value)
    if len(text) <= limit:
        return text
    return f"{text[:limit]}...(+{len(text) - limit} chars)"

def _normalize_model_name(model: str | None) -> str:
    if not isinstance(model, str):
        return ""
    return model.strip().lower()

def _resolve_model_context_tokens(model: str | None) -> int:
    override = AGENT_MODEL_CONTEXT_TOKENS or ASSISTANT_MODEL_CONTEXT_TOKENS
    if isinstance(override, int) and override > 0:
        return override
    normalized = _normalize_model_name(model)
    if not normalized:
        return 0
    if normalized in MODEL_CONTEXT_TOKENS:
        return MODEL_CONTEXT_TOKENS[normalized]
    if normalized.startswith("gpt-5.2"):
        return MODEL_CONTEXT_TOKENS.get("gpt-5.2", 0)
    return 0

def _get_encoder(model: str | None):
    normalized = _normalize_model_name(model)
    if normalized:
        try:
            return tiktoken.encoding_for_model(normalized)
        except KeyError:
            pass
    try:
        return tiktoken.get_encoding("o200k_base")
    except KeyError:
        return tiktoken.get_encoding("cl100k_base")

def _count_tokens(text: str | None, model: str | None) -> int:
    if not isinstance(text, str) or not text:
        return 0
    try:
        encoder = _get_encoder(model)
        return len(encoder.encode(text))
    except Exception:
        return max(1, len(text) // 4)

def _stringify_payload(value: Any) -> str:
    if isinstance(value, str):
        return value
    try:
        return json.dumps(value, ensure_ascii=False)
    except Exception:
        return str(value)

def _extract_text_chunks(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        chunks: list[str] = []
        for item in value:
            if isinstance(item, str):
                chunks.append(item)
                continue
            if isinstance(item, dict):
                content = item.get("content")
                if isinstance(content, str):
                    chunks.append(content)
                    continue
                if isinstance(content, list):
                    for part in content:
                        if isinstance(part, str):
                            chunks.append(part)
                        elif isinstance(part, dict):
                            text = part.get("text")
                            if isinstance(text, str):
                                chunks.append(text)
                            else:
                                chunks.append(_stringify_payload(part))
                        else:
                            chunks.append(_stringify_payload(part))
                    continue
            chunks.append(_stringify_payload(item))
        return chunks
    if isinstance(value, dict):
        content = value.get("content")
        if isinstance(content, str):
            return [content]
        if isinstance(content, list):
            chunks = []
            for part in content:
                if isinstance(part, str):
                    chunks.append(part)
                elif isinstance(part, dict):
                    text = part.get("text")
                    if isinstance(text, str):
                        chunks.append(text)
                    else:
                        chunks.append(_stringify_payload(part))
                else:
                    chunks.append(_stringify_payload(part))
            return chunks
        return [_stringify_payload(value)]
    return [str(value)]

def _calculate_context(model: str | None, instructions: str | None, input_value: Any, extra_chunks: list[str] | None = None) -> dict[str, Any]:
    chunks = []
    if isinstance(instructions, str) and instructions.strip():
        chunks.append(instructions)
    chunks.extend(_extract_text_chunks(input_value))
    if extra_chunks:
        for chunk in extra_chunks:
            if isinstance(chunk, str) and chunk:
                chunks.append(chunk)
    used_tokens = sum(_count_tokens(chunk, model) for chunk in chunks)
    max_tokens = _resolve_model_context_tokens(model)
    remaining = max(max_tokens - used_tokens, 0) if max_tokens else 0
    remaining_ratio = remaining / max_tokens if max_tokens else 0
    return {
        "maxTokens": max_tokens,
        "usedTokens": used_tokens,
        "remainingTokens": remaining,
        "remainingRatio": remaining_ratio,
    }


@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}

@app.get("/prompt", response_model=PromptResponse)
async def get_prompt() -> PromptResponse:
    return PromptResponse(prompt=_load_prompt_text())

@app.post("/prompt", response_model=PromptResponse)
async def update_prompt(req: PromptUpdateRequest) -> PromptResponse:
    try:
        prompt = _save_prompt_text(req.prompt)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return PromptResponse(prompt=prompt)

@app.get("/prompt/ui", response_class=HTMLResponse)
async def prompt_ui() -> HTMLResponse:
    prompt = html.escape(_load_prompt_text())
    page = f"""
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>Raven Prompt Editor</title>
  <style>
    :root {{
      color-scheme: light;
    }}
    body {{
      margin: 0;
      font-family: "IBM Plex Sans", "SF Pro Text", "Segoe UI", sans-serif;
      background: linear-gradient(135deg, #eef2f7, #f7f5f2);
      color: #1b1f2a;
    }}
    .wrap {{
      max-width: 920px;
      margin: 48px auto;
      padding: 0 20px 40px;
    }}
    .card {{
      background: #ffffff;
      border-radius: 18px;
      box-shadow: 0 24px 60px rgba(15, 23, 42, 0.12);
      border: 1px solid #e1e6ef;
      padding: 28px;
    }}
    h1 {{
      font-size: 22px;
      margin: 0 0 6px;
      letter-spacing: -0.01em;
    }}
    p {{
      margin: 0 0 18px;
      color: #4c5566;
    }}
    textarea {{
      width: 100%;
      min-height: 320px;
      resize: vertical;
      border-radius: 12px;
      border: 1px solid #d2d9e5;
      padding: 14px;
      font-size: 14px;
      font-family: "IBM Plex Mono", "SFMono-Regular", ui-monospace, monospace;
      line-height: 1.5;
      box-sizing: border-box;
      background: #f9fafc;
      color: #0f172a;
    }}
    .row {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      margin-top: 18px;
    }}
    button {{
      border: none;
      border-radius: 12px;
      padding: 10px 18px;
      font-weight: 600;
      font-size: 14px;
      cursor: pointer;
      background: #111827;
      color: #ffffff;
    }}
    button[disabled] {{
      opacity: 0.6;
      cursor: not-allowed;
    }}
    .status {{
      font-size: 13px;
      color: #64748b;
    }}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="card">
      <h1>Raven Prompt Editor</h1>
      <p>Edit the system prompt used by the agent service.</p>
      <textarea id="prompt">{prompt}</textarea>
      <div class="row">
        <span class="status" id="status">Ready.</span>
        <button id="save">Save</button>
      </div>
    </div>
  </div>
  <script>
    const statusEl = document.getElementById('status');
    const saveBtn = document.getElementById('save');
    const promptEl = document.getElementById('prompt');

    const setStatus = (text) => {{
      statusEl.textContent = text;
    }};

    saveBtn.addEventListener('click', async () => {{
      const text = promptEl.value || '';
      if (!text.trim()) {{
        setStatus('Prompt cannot be empty.');
        return;
      }}
      saveBtn.disabled = true;
      setStatus('Saving...');
      try {{
        const res = await fetch('/prompt', {{
          method: 'POST',
          headers: {{ 'content-type': 'application/json' }},
          body: JSON.stringify({{ prompt: text }}),
        }});
        const body = await res.json().catch(() => ({{}}));
        if (!res.ok) {{
          setStatus(body?.detail || 'Save failed.');
          return;
        }}
        setStatus('Saved.');
      }} catch (err) {{
        setStatus('Save failed.');
      }} finally {{
        saveBtn.disabled = false;
      }}
    }});
  </script>
</body>
</html>
"""
    return HTMLResponse(content=page)

@app.post("/context", response_model=AgentContextResponse)
async def get_context(req: AgentContextRequest) -> AgentContextResponse:
    instructions = _build_instructions(req.userName, req.instructions)
    context = _calculate_context(req.model, instructions, req.input)
    return AgentContextResponse(context=context)


@app.post("/run", response_model=AgentRunResponse)
async def run_agent(req: AgentRunRequest) -> AgentRunResponse:
    if not req.apiKey or not req.apiKey.strip():
        raise HTTPException(status_code=400, detail="openai_key_required")

    run_id = str(uuid.uuid4())
    started = time.monotonic()
    input_size = _estimate_size(req.input)
    logger.info(
        "run_start id=%s model=%s maxTurns=%s inputSize=%s mcp=%s",
        run_id,
        req.model,
        req.maxTurns,
        input_size,
        "yes" if (req.mcp and req.mcp.url) else "no",
    )
    if logger.isEnabledFor(logging.DEBUG):
        logger.debug(
            "run_context id=%s apiKey=%s baseUrl=%s timeoutMs=%s temperature=%s",
            run_id,
            _mask_secret(req.apiKey),
            req.openaiBaseUrl,
            req.openaiTimeoutMs,
            req.temperature,
        )
        logger.debug("run_input id=%s payload=%s", run_id, _safe_log_payload(req.input))

    timeout = None
    if isinstance(req.openaiTimeoutMs, int) and req.openaiTimeoutMs > 0:
        timeout = req.openaiTimeoutMs / 1000

    client = AsyncOpenAI(
        api_key=req.apiKey,
        base_url=req.openaiBaseUrl or None,
        timeout=timeout,
    )

    allowed = []
    if req.mcp and req.mcp.allowedTools:
        allowed = [name for name in req.mcp.allowedTools if isinstance(name, str) and name.strip()]

    if req.mcp and req.mcp.url:
        logger.info(
            "mcp_config id=%s url=%s sessionId=%s allowedTools=%s",
            run_id,
            req.mcp.url,
            req.mcp.sessionId,
            len(allowed),
        )
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(
                "mcp_details id=%s token=%s allowed=%s",
                run_id,
                _mask_secret(req.mcp.token),
                allowed,
            )

    instructions = _build_instructions(req.userName, req.instructions)
    if req.webSearchEnabled:
        instructions = "\n\n".join([
            instructions,
            "You can use the web_search tool to look up current information when needed.",
            "Use web_search only for up-to-date facts or when the user asks, and summarize the results.",
        ])
    instructions = "\n\n".join([
        instructions,
        "Use the doc_search tool when the user asks to read or search an uploaded document by URL. "
        "Provide the document URL in download_url and the question in search_request.",
    ])
    instructions = "\n\n".join([
        instructions,
        "Include an optional `reasoning` field with a short, high-level summary of your approach.",
        "Do not reveal chain-of-thought or internal reasoning steps.",
    ])
    tool_trace: list[dict[str, Any]] = []
    try:
        async with mcp_session_context(req.mcp, timeout) as (mcp_session, mcp_tools):
            function_tools = []
            if mcp_session:
                for tool in mcp_tools:
                    name = getattr(tool, "name", None)
                    if not isinstance(name, str) or not name:
                        continue
                    if allowed and name not in allowed:
                        continue
                    function_tools.append(
                        {
                            "type": "function",
                            "name": name,
                            "description": getattr(tool, "description", None),
                            "parameters": _normalize_tool_schema(getattr(tool, "inputSchema", None)),
                            "strict": False,
                        }
                    )
                logger.info(
                    "mcp_tools id=%s total=%s allowed=%s",
                    run_id,
                    len(mcp_tools),
                    len(function_tools),
                )

            doc_search_tool = {
                "type": "function",
                "name": "doc_search",
                "description": "Analyze a user-provided document by URL and answer a query. Uses direct reading for <=100k tokens, otherwise a quick RAG.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "download_url": {"type": "string", "description": "Direct URL to the file (supports txt, md, html, pdf, docx)."},
                        "search_request": {"type": "string", "description": "The user question about the document."},
                        "file_name": {"type": "string", "description": "Optional file name for better mime detection.", "nullable": True},
                    },
                    "required": ["download_url", "search_request", "file_name"],
                    "additionalProperties": False,
                },
                "strict": True,
            }

            web_search_tools = [{"type": "web_search"}] if req.webSearchEnabled else []
            tools_payload = [doc_search_tool] + function_tools + web_search_tools
            tools_enabled = bool(tools_payload)
            parse_kwargs: dict[str, Any] = {
                "model": req.model,
                "instructions": instructions,
                "input": req.input,
                "temperature": req.temperature,
                "tools": tools_payload if tools_enabled else None,
                "parallel_tool_calls": tools_enabled,
                "text_format": AssistantResponse,
            }
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(
                    "openai_request id=%s toolCount=%s",
                    run_id,
                    len(tools_payload),
                )
            response = await client.responses.parse(**parse_kwargs)

            tool_output_chunks: list[str] = []
            last_context = _calculate_context(req.model, instructions, req.input)

            while mcp_session:
                tool_calls = _prioritize_tool_calls(_extract_function_calls(getattr(response, "output", None)))
                if not tool_calls:
                    break

                outputs = []
                for call in tool_calls:
                    call_id = call.get("call_id") if isinstance(call, dict) else getattr(call, "call_id", None)
                    name = call.get("name") if isinstance(call, dict) else getattr(call, "name", None)
                    args_raw = call.get("arguments") if isinstance(call, dict) else getattr(call, "arguments", None)
                    if not call_id or not name:
                        continue
                    args = _parse_tool_args(args_raw)
                    if name == "doc_search":
                        mcp_token = req.mcp.token if req.mcp else None
                        ds = await _safe_doc_search(args, req.model, client, mcp_token=mcp_token)
                        payload_content = ds.get("result") if ds.get("ok") else {"error": ds.get("error")}
                        result = type(
                            "Result",
                            (),
                            {
                                "content": [{"type": "text", "text": json.dumps(payload_content, ensure_ascii=False)}],
                                "isError": not ds.get("ok"),
                            },
                        )
                    else:
                        result = await mcp_session.call_tool(name, args)
                    payload = {
                        "isError": bool(getattr(result, "isError", False)),
                        "content": json.loads(_serialize_tool_result(result) or "null"),
                    }
                    outputs.append(
                        {
                            "type": "function_call_output",
                            "call_id": call_id,
                            "output": json.dumps(payload, ensure_ascii=False),
                        }
                    )
                    serialized = _serialize_tool_result(result)
                    if serialized:
                        tool_output_chunks.append(serialized)
                        last_context = _calculate_context(req.model, req.instructions, req.input, tool_output_chunks)
                    trace_entry = {
                        "name": name,
                        "callId": call_id,
                        "arguments": args,
                        "output": payload["content"],
                        "isError": payload["isError"],
                    }
                    tool_trace.append(trace_entry)
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.debug(
                            "tool_call id=%s name=%s args=%s error=%s",
                            run_id,
                            name,
                            _safe_log_payload(args),
                            payload["isError"],
                        )

                if not outputs:
                    break

                parse_kwargs = {
                    "model": req.model,
                    "instructions": instructions,
                    "input": outputs,
                    "temperature": req.temperature,
                    "tools": tools_payload if tools_enabled else None,
                    "parallel_tool_calls": tools_enabled,
                    "text_format": AssistantResponse,
                    "previous_response_id": getattr(response, "id", None),
                }
                response = await client.responses.parse(**parse_kwargs)

            parsed = getattr(response, "output_parsed", None)
            output = ""
            reasoning = None
            if isinstance(parsed, AssistantResponse):
                output = parsed.message
                reasoning = parsed.reasoning
            elif isinstance(parsed, dict) and isinstance(parsed.get("message"), str):
                output = parsed["message"]
            if not output:
                output = _extract_response_text(response)
            output = _format_output(output).strip()
            if isinstance(reasoning, str):
                reasoning = reasoning.strip() or None
            elapsed = int((time.monotonic() - started) * 1000)
            logger.info(
                "run_done id=%s ms=%s outputSize=%s lastResponseId=%s",
                run_id,
                elapsed,
                len(output),
                getattr(response, "id", None),
            )
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug("run_output id=%s payload=%s", run_id, _safe_log_payload(output))
                logger.debug(
                    "openai_response id=%s usage=%s outputTypes=%s",
                    run_id,
                    _safe_log_payload(getattr(response, "usage", None)),
                    _summarize_output_items(getattr(response, "output", None)),
                )
            extra_chunks = tool_output_chunks[:]
            if output:
                extra_chunks.append(output)
            context = _calculate_context(req.model, instructions, req.input, extra_chunks) if extra_chunks else last_context
            if req.webSearchEnabled:
                web_search_items = _extract_web_search_items(getattr(response, "output", None))
                for _item in web_search_items:
                    tool_trace.append({"name": "web_search"})
            trace = None
            if tool_trace or reasoning:
                trace = {
                    "reasoning": reasoning,
                    "tools": tool_trace if tool_trace else None,
                }
            return AgentRunResponse(
                output=output,
                lastResponseId=getattr(response, "id", None),
                context=context,
                trace=trace,
            )
    except APIStatusError as exc:
        code = None
        message = str(exc)
        if isinstance(exc.body, dict):
            err = exc.body.get("error") or {}
            if isinstance(err, dict):
                code = err.get("code") or err.get("type")
                message = err.get("message") or message
        elapsed = int((time.monotonic() - started) * 1000)
        logger.error(
            "run_error id=%s ms=%s status=%s code=%s message=%s",
            run_id,
            elapsed,
            exc.status_code,
            code,
            message,
        )
        raise HTTPException(
            status_code=exc.status_code or 500,
            detail={"error": code or "openai_error", "message": message},
        ) from exc
